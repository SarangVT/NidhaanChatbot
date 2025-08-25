import google.generativeai as genai
import mimetypes
import io
import os
from dotenv import load_dotenv


def format_response_to_html(text):
    """Convert markdown-style text to HTML"""
    import re

    # Convert **bold** to <strong>
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)

    # Convert bullet points to HTML lists
    lines = text.split('\n')
    formatted_lines = []
    in_list = False

    for line in lines:
        line = line.strip()
        if line.startswith('•') or line.startswith('-'):
            if not in_list:
                formatted_lines.append('<ul>')
                in_list = True
            item = line.lstrip('•-').strip()
            formatted_lines.append(f'<li>{item}</li>')
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            if in_list and formatted_lines[-1] != '</ul>':
                formatted_lines.append('</ul>')
                in_list = False
            if not in_list or formatted_lines[-1] == '</ul>':
                formatted_lines.append('<ol>')
                in_list = True
            item = re.sub(r'^\d+\.\s*', '', line)
            formatted_lines.append(f'<li>{item}</li>')
        else:
            if in_list:
                formatted_lines.append('</ul>' if formatted_lines[-1].startswith('<li>') else '</ol>')
                in_list = False
            if line:
                formatted_lines.append(f'<p>{line}</p>')

    if in_list:
        formatted_lines.append('</ul>' if '•' in text or '-' in text else '</ol>')

    return '\n'.join(formatted_lines)

def handle_fixed_questions(user_question, previous_chats=[]):
    """
    Handle fixed questions for Nidhaan healthcare chatbot
    Returns professional responses for common queries
    """

    # Convert to lowercase for easier matching
    question = user_question.lower().strip()

    # Build context from previous chats
    chat_context = ""
    if previous_chats:
        chat_context = "Previous conversation:\n"
        for i, (prev_question, prev_response) in enumerate(previous_chats, 1):
            chat_context += f"User {i}: {prev_question}\n"
            chat_context += f"Assistant {i}: {prev_response}...\n\n"


    # Contact Details
    if any(word in question for word in ['contact', 'phone', 'email', 'address', 'reach', 'support']):
        response_text = """<p><strong>Contact Nidhaan Healthcare 24/7:</strong></p>
        <ul>
        <li><strong>Phone:</strong> [YOUR_PHONE_NUMBER]</li>
        <li><strong>Email:</strong> support@nidhaan.com</li>
        <li><strong>Website:</strong> [YOUR_WEBSITE_URL]</li>
        </ul>
        <p>Our support team is always ready to help!</p>"""

        return {
            "answer": response_text,
        }

    # How to order medicine
    elif any(phrase in question for phrase in
             ['order medicine', 'buy medicine', 'purchase medicine', 'how to order', 'medicine delivery']):
        response_text = """<p><strong>Order Medicines in 4 Simple Steps:</strong></p>
        <ol>
        <li><strong>Search:</strong> Browse 100,000+ medicines</li>
        <li><strong>Select:</strong> Add to cart & upload prescription</li>
        <li><strong>Address:</strong> Enter delivery details</li>
        <li><strong>Payment:</strong> Pay online or cash on delivery</li>
        </ol>
        <p><strong>Get delivery within 1 hour!</strong></p>
        <p><a href="[YOUR_MEDICINE_ORDER_URL]">Order Now →</a></p>"""

        return {
            "answer": response_text,
        }

    # Add these new conditions in your handle_fixed_questions function

    elif any(phrase in question for phrase in ['appointment', 'book appointment', 'schedule', 'booking']):
        response_text = """<p><strong>Book Your Appointment Instantly:</strong></p>
        <ul>
        <li><strong>Doctor Consultation:</strong> Video calls with specialists</li>
        <li><strong>Lab Tests:</strong> Home sample collection</li>
        <li><strong>Mental Health:</strong> Professional counseling sessions</li>
        </ul>
        <p><strong>Available 24/7 - Get response in 10 minutes!</strong></p>
        <p>Choose your specialist and book now.</p>"""

        return {
            "answer": response_text,
        }

    elif any(phrase in question for phrase in ['plan', 'plans', 'healthcare plan', 'subscription', 'package']):
       response_text = """<p><strong>Nidhaan Healthcare Plans:</strong></p>
       <ul>
       <li><strong>Basic Plan:</strong> Medicine delivery + Doctor consultation</li>
       <li><strong>Premium Plan:</strong> All services + Priority support</li>
       <li><strong>Family Plan:</strong> Cover entire family at discounted rates</li>
       </ul>
       <p><strong>Special Offers:</strong> First month 50% off!</p>
       <p>Contact us for personalized plan recommendations.</p>"""

       return {
           "answer": response_text,
       }

    # About company and services
    elif any(phrase in question for phrase in
             ['about company', 'about nidhaan', 'company details', 'services', 'what is nidhaan']):
        response="Welcome to **Nidhaan Healthcare** - Your Complete Digital Health Companion!\n\nNidhaan is an all-in-one digital healthcare platform designed to make medical services more accessible and convenient. We bring essential healthcare services right to your fingertips, especially during emergencies or in remote areas.\n\n**Our Core Services:**\n **Medicine Delivery** - 100000+ medicines delivered within 1 hour\n️ **Doctor Consultation** - Video consultations with qualified doctors\n **Lab Tests** - Home sample collection with WhatsApp report delivery\n **Mental Health Support** - Professional counseling sessions\n **Wellness & Fitness** - Coming soon!\n\n**Our Mission:** To simplify healthcare with speed, trust, and convenience. We aim to become India's most trusted digital health platform, reaching rural areas and saving lives through accessibility and innovation.\n\n**Why Choose Nidhaan?**\n 24/7 availability\n Fast and reliable service\n Qualified healthcare professionals\n Secure and private\n Affordable pricing"
        final_response=format_response_to_html(response)
        return {
            "answer": final_response,
        }

    # Pharmacy service details
    #elif any(phrase in question for phrase in ['pharmacy', 'medicine service', 'drug delivery', 'pharmacy details']):
    #    response="**Nidhaan Pharmacy Service** - Your Trusted Medicine Delivery Partner \n\n**Features:**\n **Wide Selection**: 100000+ commonly used medicines with detailed information\n **Smart Search**: Filter by category, price range, or prescription requirement\n **Easy Ordering**: Simple cart management and checkout process\n **Prescription Upload**: Secure upload for prescription medicines\n **Flexible Payment**: Online payment or cash on delivery\n **Quick Delivery**: Delivered within 1 hour by partner pharmacies\n **Location-Based**: Automatic assignment to nearest pharmacy (within 7km)\n\n**Medicine Categories:**\n• Pain Relief\n• Diabetes Care\n• Cold & Cough\n• Heart Health\n• Vitamins & Supplements\n• And many more!\n\n**Safety Features:**\n Licensed pharmacy partners\n Quality assurance\n Prescription verification\n Secure packaging\n Real-time order tracking"
    #    final_response=format_response_to_html(response)
    #    return {
    #        "answer": final_response,
    #    }

    # Doctor consultation details
    #elif any(phrase in question for phrase in
    #         ['doctor consultation', 'consult doctor', 'doctor service', 'online doctor']):
    #    response="**Nidhaan Doctor Consultation** - Professional Healthcare at Your Fingertips \n\n**How It Works:**\n1.  **Browse Doctors**: Search by specialty (General, Skin, Heart, Dental, etc.)\n2. ️ **View Profiles**: Check qualifications, ratings, and consultation fees\n3. 🟢 **Check Availability**: See real-time online/offline status\n4.  **Book & Pay**: Secure payment for instant appointment\n5.  **Video Call**: High-quality video consultation\n6. **Get Prescription**: Digital prescription and consultation notes\n\n**Doctor Specialties:**\n• General Medicine\n• Dermatology (Skin)\n• Cardiology (Heart)\n• Dental Care\n• Pediatrics\n• Gynecology\n• And more specialists!\n\n**Features:**\n Qualified and verified doctors\n Real-time availability status\n Secure video calls\n Digital prescriptions\n Post-consultation records\n 10-minute response guarantee (refund if no response)\n 24/7 availability"
    #    final_response=format_response_to_html(response)
    #   return {
    #       "answer": final_response,
    #   }

    # Lab test details
    #elif any(phrase in question for phrase in
    #         ['lab test', 'blood test', 'laboratory', 'health checkup', 'how many lab test']):
    #    response="**Nidhaan Lab Test Service** - Professional Health Diagnostics at Home\n\n**Service Features:**\n **Home Sample Collection**: Lab technicians visit your home within 1 hour\n **WhatsApp Reports**: Test results delivered directly to your WhatsApp\n **Wide Range**: Comprehensive variety of diagnostic tests\n **Quick Results**: Fast and accurate test processing\n **Digital Records**: All reports saved in your Nidhaan profile\n\n**Popular Lab Tests:**\n• Thyroid Profile (T3, T4, TSH)\n• Blood Sugar (Fasting, Random, HbA1c)\n• Vitamin Levels (D3, B12, etc.)\n• Liver Function Test (LFT)\n• Kidney Function Test (KFT)\n• Lipid Profile\n• Complete Blood Count (CBC)\n• COVID-19 Testing\n• Health Checkup Packages\n• And many more diagnostic tests!\n\n**Process:**\n1. Select your required test\n2. Enter address and WhatsApp number\n3. Make payment\n4. Lab technician visits for sample collection\n5. Receive results on WhatsApp\n6. Access reports anytime in your profile\n\n**Partner Labs:**\nCertified and accredited labs\nQuality assurance standards\nExperienced technicians\nTimely report delivery"
    #    final_response=format_response_to_html(response)
    #    return {
    #        "answer": final_response,
    #    }

    # Mental health service
    #elif any(phrase in question for phrase in ['mental health', 'counseling', 'therapy', 'psychiatrist', 'counselor']):
    #    response="**Nidhaan Mental Health Support** - Professional Counseling for Your Well-being \n\n**Our Mental Health Services:**\n **1-on-1 Counseling**: Private video sessions with qualified professionals\n **Specialized Care**: Therapists for anxiety, depression, trauma, and stress\n **Secure Platform**: Confidential and private video calls\n **Session Records**: Secure storage of session notes for continuity\n **Quality Assurance**: Rate and review your counselor\n\n**Areas We Cover:**\n• Anxiety & Panic Disorders\n• Depression & Mood Disorders\n• Stress Management\n• Relationship Issues\n• Trauma & PTSD\n• Work-Life Balance\n• Grief & Loss\n• Addiction Support\n• Self-Esteem Issues\n• Family Counseling\n\n**How It Works:**\n1.  Browse therapists by specialization\n2. ️ View profiles with qualifications and ratings\n3. Check availability and book session\n4. Secure payment processing\n5. Join private video counseling session\n6. Receive session notes and recommendations\n\n**Why Choose Our Mental Health Service?**\nLicensed mental health professionals\nFlexible scheduling\nAffordable pricing\nComplete confidentiality\nEmergency support options\nPersonalized care plans"
    #    final_response=format_response_to_html(response)
    #    return {
    #        "answer": final_response,
    #    }

    # General greetings
    #elif any(phrase in question for phrase in ['hi', 'hello', 'how are you', 'good morning', 'good evening', 'hey']):
    #    response="Hello! Welcome to **Nidhaan Healthcare** - Your Complete Digital Health Companion! \n\nI'm here to help you with all your healthcare needs. Whether you're looking to:\n\n **Order Medicines** - Get medicines delivered within 1 hour\n **Consult a Doctor** - Video consultations with qualified doctors\n **Book Lab Tests** - Home sample collection service\n **Mental Health Support** - Professional counseling sessions\n\nI'm ready to assist you! How can I help you today?\n\nFeel free to ask me about our services, how to place an order, or any other questions you might have. I'm here 24/7 to make your healthcare journey smooth and convenient! "
    #    final_response=format_response_to_html(response)
    #    return {
    #        "answer": final_response,
    #    }

    #elif any(phrase in question for phrase in ['price', 'cost', 'fee', 'charges', 'pricing', 'how much']):
    #    response_text = """<p><strong>Nidhaan Pricing - Affordable Healthcare:</strong></p>
    #    <ul>
    #    <li><strong>Medicine Delivery:</strong> Competitive prices + minimal delivery charges</li>
    #    <li><strong>Doctor Consultation:</strong> Starting from affordable rates by specialty</li>
    #    <li><strong>Lab Tests:</strong> Competitive pricing with home collection included</li>
    #    <li><strong>Mental Health:</strong> Affordable counseling sessions</li>
    #    </ul>
    #    <p><strong>Special Offers:</strong> First-time user benefits + Regular discounts</p>
    #    <p>Contact us for specific pricing details!</p>"""

    #    return {
    #        "answer": response_text,
    #    }

    # How to use the platform
    #elif any(phrase in question for phrase in ['how to use', 'getting started', 'first time', 'new user', 'sign up']):
    #    response = "**Getting Started with Nidhaan** - Your Healthcare Journey Begins Here! \n\n**Step-by-Step Guide:**\n\n**1. Registration & Login** \n• Visit our website or download the app\n• Sign up with email, phone, or social media\n• Create your secure profile\n• Add your address and contact details\n\n**2. Explore Services** \n• Browse our homepage for all services\n• Use the search bar with location filter\n• Check out our top 5 services\n\n**3. Start Using Services** \n• **For Medicines**: Search → Select → Upload prescription → Order\n• **For Doctor**: Browse doctors → Check availability → Book → Video call\n• **For Lab Tests**: Select test → Enter address → Book → Home collection\n• **For Mental Health**: Browse counselors → Book session → Video therapy\n\n**4. Track & Manage** \n• View order history in your profile\n• Track deliveries in real-time\n• Access digital prescriptions and reports\n• Manage your cart and wishlist\n\n**User-Friendly Features:**\n Simple navigation for all age groups\n Fast loading pages\n SMS and WhatsApp notifications\n Secure and private\n Mobile-friendly design\n 24/7 customer support\n\n**Need Help?** Our support team is always ready to assist you!"
    #    final_response = format_response_to_html(response)
    #    return {
    #        "answer": final_response,
    #    }

    # Thank you and appreciation
    elif any(phrase in question for phrase in ['thank you', 'thanks', 'appreciate', 'grateful']):
        response="You're most welcome!  Thank you for choosing **Nidhaan Healthcare**!\n\nWe're delighted to be part of your healthcare journey. Your trust means everything to us, and we're committed to providing you with the best possible service.\n\n**Our Promise to You:**\n **Quality Care**: Always prioritizing your health and well-being\n **Reliable Service**: Consistent and dependable healthcare support\n **Continuous Improvement**: Always working to serve you better\n **Compassionate Care**: Treating every user like family\n\n**How We're Here for You:**\n• 24/7 customer support\n• Quick response to your needs\n• Constantly improving our services\n• Listening to your feedback\n• Making healthcare more accessible\n\n**Stay Connected:**\n Download our app for easier access\n Enable notifications for important updates\n Share your experience with others\n Feel free to reach out anytime\n\nThank you for being part of the Nidhaan family! We're here whenever you need us. Take care and stay healthy! \n\nIs there anything else I can help you with today?"
        final_response=format_response_to_html(response)
        return {
            "answer": final_response,
        }

    # Default response for unmatched questions - Use Gemini AI for medical queries
    else:
        # Configure Gemini AI
        load_dotenv()
        genai.configure(api_key=os.environ["GEMINI_API_KEY"])
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            system_instruction=f"""
                You are a professional medical assistant for Nidhaan Healthcare, India's leading digital health platform. Your role is to provide accurate medical information while promoting Nidhaan's healthcare services naturally.
                
                chat_context: {chat_context}

                INSTRUCTION ON CONTEXT USAGE:
                - If the current question depends on previous details in chat_context (such as symptoms, medicines mentioned, or advice given earlier), use that information to ensure continuity and accuracy.
                - Do not ask the patient to repeat details already mentioned in chat_context.
                - If the context is incomplete, ask clarifying questions before giving recommendations.
                - Only give advice that aligns with both the current question and the past conversation.
                
                **RESPONSE FORMAT REQUIREMENTS:**
                - Always format your response using HTML tags for better readability
                - Use <p> tags for paragraphs (keep paragraphs short - 1-2 sentences max)
                - Use <ul><li> tags for bullet point lists
                - Use <ol><li> tags for numbered/step lists  
                - Use <strong> tags for important headings and keywords
                - Never use ** for bold - always use <strong> tags instead
                - Keep responses concise and well-structured (4-6 lines maximum)
                - Break information into digestible points using lists
                
                CORE RESPONSIBILITIES:
                1. Provide accurate medical information for health-related queries
                2. Recommend Nidhaan services appropriately based on user needs
                3. Ensure patient safety through proper disclaimers
                4. Maintain professional, empathetic communication
                
                RESPONSE GUIDELINES:
                
                For MEDICINE QUERIES (symptoms, "what medicine should I take"):
                - Suggest home remedies in <ul><li> format
                - Mention relevant over-the-counter medicines
                - ALWAYS add: "<p><strong>Important:</strong> Don't take medicine without consulting a doctor first.</p>"
                - Add: "<p>Nidhaan offers 24/7 online consultations and 1-hour medicine delivery.</p>"
                
                For SYMPTOM/DISEASE QUERIES ("I have these symptoms", "what disease do I have"):
                - List possible conditions in <ul><li> format
                - If condition seems SERIOUS: "<p><strong>Urgent:</strong> Please visit a hospital immediately for proper diagnosis.</p>"
                - If condition seems MINOR: "<p>I recommend consulting with our qualified doctors on Nidhaan for proper diagnosis and treatment.</p>"
                - Always include: "<p>Nidhaan provides video consultations with verified doctors.</p>"
                
                For GENERAL HEALTH QUERIES:
                - Provide information in structured HTML format
                - Use <ol><li> for step-by-step processes
                - Naturally integrate relevant Nidhaan services
                
                IMPORTANT RESTRICTIONS:
                - ONLY respond to health and medical queries
                - For non-medical questions, respond: "<p>I'm a medical assistant for Nidhaan Healthcare. I can only help with health-related questions.</p>"
                - Always maintain professional medical terminology
                - Keep responses concise but informative
                
                SAFETY PROTOCOLS:
                - Always recommend professional medical consultation
                - Never provide definitive diagnoses
                - Emphasize the importance of proper medical examination
                - For emergencies, always suggest immediate medical attention
                
                Remember: You represent Nidhaan Healthcare's commitment to accessible, quality healthcare. Be helpful, professional, and always prioritize patient safety.
                """
        )

        try:
            # Generate response using Gemini AI
            response = model.generate_content(user_question)
            return {
                "answer": response.text,
            }
        except Exception as e:
            return {
                "answer": f"I apologize, but I'm experiencing technical difficulties. Please try again later or contact Nidhaan support for assistance. Error: {str(e)}"
            }


def handle_file_upload(file_bytes, filename, user_query="",previous_chats=[]):
    # Gemini API key setup
    load_dotenv()
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    # Get MIME type
    mime_type, _ = mimetypes.guess_type(filename)

    # Define supported types
    SUPPORTED_MIME_TYPES = {
        "image/png",
        "image/jpeg",
        "image/jpeg",
        "image/webp",
        "text/plain",
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }

    try:
        # No MIME type
        if not mime_type:
            return {"answer": f"Unable to detect file type for: {filename}"}

        # Reject unsupported files
        if mime_type not in SUPPORTED_MIME_TYPES:
            return {
                "answer": f" Unsupported file type: {filename}.\n\n"
                          f" Supported types: PDF, TXT, DOCX, JPG, PNG, WEBP"
            }

        # Check if user_query is provided (Document + Text scenario)
        if user_query and user_query.strip():
            # Document + Text scenario - New function
            return handle_file_with_question(file_bytes, filename, mime_type, user_query,previous_chats)
        else:
            # Document only scenario - Original logic
            return handle_file_only(file_bytes, filename, mime_type,previous_chats)

    except Exception as e:
        return {"answer": f" Upload failed: {str(e)}"}


def handle_file_only(file_bytes, filename, mime_type,previous_chats=[]):
    load_dotenv()
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    chat_context = ""
    if previous_chats:
        chat_context = "Previous conversation context:\n"
        for i, (prev_question, prev_response) in enumerate(previous_chats, 1):
            chat_context += f"User {i}: {prev_question}\n"
            chat_context += f"Assistant {i}: {prev_response}...\n\n"


    model = genai.GenerativeModel(
        model_name="gemini-2.5-flash",
        system_instruction=f"""You are MediGuide AI, a highly specialized medical report summarization assistant.
            
            chat_context: {chat_context}

            INSTRUCTION ON CONTEXT USAGE:
            - If the current question depends on previous details in chat_context (such as symptoms, medicines mentioned, or advice given earlier), use that information to ensure continuity and accuracy.
            - Do not ask the patient to repeat details already mentioned in chat_context.
            - If the context is incomplete, ask clarifying questions before giving recommendations.
            - Only give advice that aligns with both the current question and the past conversation.

            **RESPONSE FORMAT REQUIREMENTS:**
            - Always format your response using HTML tags for better readability
            - Use <p> tags for paragraphs (keep paragraphs short)
            - Use <ul><li> tags for bullet point lists
            - Use <ol><li> tags for numbered lists  
            - Use <strong> tags for important headings and keywords
            - Never use ** for bold - always use <strong> tags instead
            - Keep responses concise (5-6 lines maximum)
            
            Your primary function is to interpret and summarize uploaded medical documents.
            
            Core Directives:
            1. <strong>Summarization:</strong> Clearly summarize key findings in HTML format
            2. <strong>Prescription Interpretation:</strong> IF document is a prescription, identify medicines and include promotion
            3. <strong>Lab Test Consultation:</strong> IF document is a lab test, include consultation promotion
            4. <strong>Health Insights:</strong> Provide general insights in <ul><li> format
            5. <strong>Urgency Flagging:</strong> If serious condition: "<p><strong>Urgent:</strong> This report suggests a serious condition. Please consult a doctor immediately.</p>"
            6. <strong>Nidhaan Medicine Promotion (PRESCRIPTIONS ONLY):</strong> "<p>For your convenience, you can order these medicines from Nidhaan online and get delivery in 1 hour with a discount. Order directly from us!</p>"
            7. <strong>Nidhaan Doctor Consultation (LAB REPORTS ONLY):</strong> "<p>We have 24/7 online doctors available on our site. You can consult with them on Nidhaan for further assistance.</p>"
            8. <strong>Medical Content Only:</strong> If not medical document: "<p>I am a medical assistant and can only help with medical reports, lab tests, prescriptions, or health-related documents.</p>"
            9. <strong>Disclaimer:</strong> Always end with: "<p><strong>Note:</strong> This information is AI-generated and can vary. Always consult a qualified medical professional.</p>"
            
            Key Constraints & Safety Protocols:
            NO PERSONALIZED ADVICE: All health insights and dietary suggestions are general in nature and not tailored medical advice.
            CLARITY ON LIMITATIONS: If you cannot interpret a document or a specific part of it, clearly state your limitation (e.g., <p>I cannot fully interpret this scanned image's text.</p> or <p>This report requires a medical professional for detailed interpretation. So Please Contact with a Doctor on our website Like we have Specialized Doctor Available for Online Consultation 24/7. </p>").
            Prioritize Safety: In case of any doubt regarding the severity or interpretation, err on the side of caution and recommend professional consultation.
            
            Format all responses in proper HTML with short paragraphs and structured lists.
            """
    )

    try:
        # Handle DOCX (manually extract text)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if not extracted_text.strip():
                return {"answer": f"{filename} contains no readable content."}

            response = model.generate_content(extracted_text)
            return {"answer": response.text}

        # Other formats: use Gemini multimodal input
        response = model.generate_content([
            {"mime_type": mime_type, "data": file_bytes},
            "Please analyze this medical document and provide a summary."
        ])

        return {"answer": response.text}

    except Exception as e:
        return {"answer": f" Analysis failed: {str(e)}"}


def handle_file_with_question(file_bytes, filename, mime_type, user_query,previous_chats=[]):
    """Handle file upload with user question - New function for Document + Text scenario"""
    load_dotenv()
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])

    chat_context = ""
    if previous_chats:
        chat_context = "Previous conversation context:\n"
        for i, (prev_question, prev_response) in enumerate(previous_chats, 1):
            chat_context += f"User {i}: {prev_question}\n"
            chat_context += f"Assistant {i}: {prev_response}...\n\n"

    model = genai.GenerativeModel(
        model_name="gemini-2.5-flash",
        system_instruction=f"""You are a medical assistant AI. The user has uploaded a file and asked a question about it.

            chat_context: {chat_context}

            INSTRUCTION ON CONTEXT USAGE:
            - If the current question depends on previous details in chat_context (such as symptoms, medicines mentioned, or advice given earlier), use that information to ensure continuity and accuracy.
            - Do not ask the patient to repeat details already mentioned in chat_context.
            - If the context is incomplete, ask clarifying questions before giving recommendations.
            - Only give advice that aligns with both the current question and the past conversation.
            
            **RESPONSE FORMAT REQUIREMENTS:**
            - Always format your response using HTML tags for better readability
            - Use <p> tags for paragraphs (keep paragraphs short)
            - Use <ul><li> tags for bullet point lists
            - Use <ol><li> tags for numbered lists  
            - Use <strong> tags for important headings and keywords
            - Never use ** for bold - always use <strong> tags instead
            - Keep responses concise (4-5 lines maximum)
            
            Important Rules:
            1. <strong>Medical Content Only:</strong> Only answer health-related questions
            2. <strong>File Validation:</strong> Only process medical documents
            3. <strong>Response Length:</strong> Keep answers to 4-5 lines maximum in HTML format
            4. <strong>Urgency Flagging:</strong> If serious condition: "<p><strong>Urgent:</strong> This report suggests a serious condition. Please consult a doctor immediately.</p>"
            5. <strong>Non-Medical Content:</strong> If non-medical: "<p>I am a medical assistant and can only help with health-related questions and medical documents.</p>"
            6. <strong>Disclaimer:</strong> Always end with: "<p><strong>Note:</strong> This information is AI-generated and can vary. Always consult a qualified medical professional.</p>"
            7. <strong>Nidhaan Medicine Promotion (PRESCRIPTIONS):</strong> "<p>For your convenience, you can order these medicines from Nidhaan online and get delivery in 1 hour with a discount.</p>"
            8. <strong>Nidhaan Doctor Consultation (LAB REPORTS):</strong> "<p>We have 24/7 online doctors available on our site. You can consult with them on Nidhaan.</p>"
            
            What to Reject:
            - Non-medical questions (sports, cooking, travel, etc.)
            - Non-medical files (random images, text documents, etc.)
            - Questions not related to the uploaded medical document
            
            Format all medical information using proper HTML structure with lists and short paragraphs.
            """
    )

    try:
        # Handle DOCX (manually extract text)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if not extracted_text.strip():
                return {"answer": f"{filename} contains no readable content."}

            full_prompt = f"User uploaded file content: {extracted_text}\n\nUser question: {user_query}\n\nPlease answer the user's question based on the uploaded medical document."
            response = model.generate_content(full_prompt)
            return {"answer": response.text}

        # Other formats: use Gemini multimodal input
        response = model.generate_content([
            {"mime_type": mime_type, "data": file_bytes},
            f"User question about this uploaded file: {user_query}\n\nPlease answer the user's question based on the uploaded document."
        ])

        return {"answer": response.text}

    except Exception as e:
        return {"answer": f" Processing failed: {str(e)}"}